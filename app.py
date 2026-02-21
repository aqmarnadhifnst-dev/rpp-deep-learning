import streamlit as st
from docx import Document
from io import BytesIO

# --- KONFIGURASI HALAMAN ---
st.set_page_config(page_title="Generator RPP - MAS Darul Ikhlas", page_icon="📝")

st.title("📝 Generator Modul Ajar (Deep Learning)")
st.subheader("Oleh: Hanapi - MAS Darul Ikhlas")

# --- INPUT DATA ---
with st.sidebar:
    st.header("Identitas Modul")
    nama_guru = st.text_input("Nama Guru", "Hanapi, S.Pd.")
    mapel = st.text_input("Mata Pelajaran")
    kelas = st.selectbox("Kelas", ["X", "XI", "XII"])
    materi = st.text_area("Topik Materi Utama")

st.info("Isi detail pembelajaran di bawah ini untuk menghasilkan file Word.")

tujuan = st.text_area("Tujuan Pembelajaran (Deep Learning)")
langkah_langkah = st.text_area("Langkah Pembelajaran (Pembukaan, Inti, Penutup)")
asesmen = st.text_area("Metode Asesmen")

# --- LOGIKA PEMBUATAN FILE WORD ---
def generate_docx():
    doc = Document()
    doc.add_heading('MODUL AJAR / RPP PEMBELAJARAN MENDALAM', 0)
    
    # Header Institusi
    p = doc.add_paragraph()
    p.add_run(f"MAS DARUL IKHLAS\n").bold = True
    p.add_run(f"Penyusun: {nama_guru}\n")
    p.add_run(f"Mapel: {mapel} | Kelas: {kelas}")
    
    doc.add_heading('A. Tujuan Pembelajaran', level=1)
    doc.add_paragraph(tujuan)
    
    doc.add_heading('B. Langkah-Langkah (Deep Learning)', level=1)
    doc.add_paragraph(langkah_langkah)
    
    doc.add_heading('C. Asesmen', level=1)
    doc.add_paragraph(asesmen)
    
    doc.add_paragraph("\nMengetahui,\nKepala Madrasah\n\n\n(____________________)")
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- TOMBOL DOWNLOAD ---
if st.button("Generate File Word"):
    if mapel and materi:
        doc_file = generate_docx()
        st.download_button(
            label="⬇️ Download RPP (Word)",
            data=doc_file,
            file_name=f"RPP_{mapel}_{kelas}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.error("Mohon isi minimal Nama Mapel dan Materi!")
